"""
translate_manual.py — แปลคู่มือเครื่องจักร (สแกน/ภาพ) EN -> ไทย ด้วย gpt-4.1 vision
ออกผลเป็น Word (.docx) + PDF (.pdf)

ออกแบบมาสำหรับเล่มหนา ~1000 หน้า:
  * แปลทีละหน้าด้วย gpt-4.1 vision (อ่าน OCR + แปล ในขั้นตอนเดียว ไม่ต้อง tesseract)
  * checkpoint รายหน้า -> รันซ้ำ/รันต่อได้ ไม่จ่ายค่าแปลซ้ำ
  * ทำหลายหน้าพร้อมกัน (--workers)
  * markdown -> docx (renderer ในตัว: หัวข้อ/ตาราง/ตัวหนา/ภาพ)
  * docx -> pdf ด้วย LibreOffice headless

ตัวอย่างใช้งาน:
  python translate_manual.py "C:\\Users\\...\\Test2.pdf" --out gl44_safety
  python translate_manual.py manual.pdf --pages 1-50 --workers 4 --dpi 200

ไม่ยุ่งกับ backend service / database — เป็นเครื่องมือ stand-alone อ่าน PDF เขียนไฟล์ออกเท่านั้น
"""
from __future__ import annotations

import argparse
import base64
import os
import re
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz  # PyMuPDF
from dotenv import load_dotenv
from openai import OpenAI

# ---------------------------------------------------------------- config

HERE = Path(__file__).resolve().parent
load_dotenv(HERE / ".env")

MODEL = os.getenv("LLM_MODEL", "gpt-4.1")
TEXT_LAYER_MIN = 60   # ตัวอักษรขั้นต่ำในหน้า ถึงจะถือว่าเป็น "หน้าดิจิทัล" (ใช้ไฮบริด)
# gpt-4.1 ราคาโดยประมาณ (USD ต่อ 1M tokens) — ใช้ประเมินต้นทุนเท่านั้น
PRICE_IN = 2.0
PRICE_OUT = 8.0
USD_TO_THB = 36.0

THAI_FONT = "Sarabun"  # ฟอนต์มาตรฐานเอกสารไทย อ่านง่าย สวย (ติดตั้งจาก Google Fonts)
FONT_DIR = HERE / "fonts"  # ไฟล์ Sarabun-*.ttf สำหรับฝังลง .docx (OFL อนุญาตให้ฝัง)

SOFFICE_CANDIDATES = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice",
    "/usr/bin/libreoffice",
    "/opt/libreoffice/program/soffice",
]


def _find_soffice() -> str | None:
    import shutil
    for name in ("soffice", "libreoffice"):
        p = shutil.which(name)
        if p:
            return p
    return next((p for p in SOFFICE_CANDIDATES if Path(p).exists()), None)

# ---------------------------------------------------------------- glossary

# ศัพท์เทคนิคล็อกให้แปลเหมือนกันทั้งเล่ม (gpt อ้างอิงตอนแปล)
GLOSSARY = """\
plate cylinder = โมแม่พิมพ์ (plate cylinder)
blanket cylinder = โมผ้ายาง (blanket cylinder)
impression cylinder = ลูกโมพิมพ์ (impression cylinder)
dampening / fountain roller = ลูกกลิ้งน้ำยา (dampening fountain roller)
ink fountain = รางหมึก (ink fountain)
inking section = ส่วนหมึก (inking section)
feeder = ชุดป้อนกระดาษ (feeder)
delivery = ชุดส่งกระดาษ (delivery)
inching = เดินเครื่องทีละน้อย (inching)
safety bar = แถบนิรภัย (safety bar)
safety cover = ฝาครอบนิรภัย (safety cover)
emergency stop = หยุดฉุกเฉิน (emergency stop)
register unit = ชุดเรจิสเตอร์ (register unit)
"""

SYSTEM_PROMPT = "คุณเป็นนักแปลคู่มือเครื่องจักรการพิมพ์ Komori มืออาชีพ แปลอังกฤษเป็นไทยอย่างถูกต้องและครบถ้วน"

PAGE_PROMPT = f"""\
แปลหน้าคู่มือเครื่องพิมพ์นี้จากภาษาอังกฤษเป็นภาษาไทย ออกเป็น Markdown ตามกฎ:

1. ถอดและแปลข้อความ "ทุกส่วน" บนหน้า ห้ามย่อ ห้ามข้าม
2. คงโครงสร้างเดิม: หัวข้อ, เลข section (เช่น 1-4), เลขรายการ, ตาราง (ใช้ Markdown table)
3. ศัพท์เทคนิคใช้ตามอภิธานนี้ ให้ตรงกันทั้งเล่ม:
{GLOSSARY}
4. ป้าย/ปุ่มจริงบนเครื่อง และคำเตือนระดับหัวข้อ คงภาษาอังกฤษ: WARNING, CAUTION, DANGER, TUNE UP, OVERLOAD, ON/OFF
5. รูปภาพ/ไดอะแกรม: เขียนเป็น `[ภาพ: คำอธิบายสั้นๆ ว่ารูปแสดงอะไร]` — ห้ามแปลตัวเลข callout ในรูป ให้คงภาพต้นฉบับไว้ (จะแปะภาพให้ทีหลัง)
6. **ข้ามแบนเนอร์หัว/ท้ายกระดาษที่ซ้ำทุกหน้า** (ชื่อบริษัท ชื่อคู่มือ เลขหน้า) — ไม่ต้องแปลหรือใส่ในผลลัพธ์
7. อย่าใส่คำอธิบายของคุณเอง ตอบเฉพาะเนื้อหาที่แปลแล้ว ไม่ต้องครอบด้วย ```
8. ส่วนที่จาง/เลือน/ถูกเซ็นเซอร์/ว่างเปล่า อ่านไม่ออก: เขียนแค่ `[ส่วนนี้จาง อ่านไม่ออก — ต้องสแกนใหม่]` ครั้งเดียวแล้วข้ามไป ห้ามเดา ห้ามแต่งเติม
9. ห้ามใส่ตัวอักษร/ช่องว่าง/`&nbsp;`/จุด ซ้ำๆ เพื่อจัดช่องว่างเด็ดขาด ถ้าไม่มีเนื้อหาก็ไม่ต้องพิมพ์อะไร\
"""

VERIFY_SYSTEM = (
    "คุณเป็นผู้ตรวจสอบคุณภาพการแปลคู่มือเครื่องจักรนิรภัย ละเอียดและเข้มงวด "
    "ตอบเป็น JSON เท่านั้น"
)

# รอบตรวจทาน — เน้นจับ "สลับแถว" จากตารางเซลล์ผสาน (merged cell) ซึ่งเป็นจุดอ่อนหลักของ vision
VERIFY_PROMPT = """\
นี่คือภาพหน้าคู่มือต้นฉบับ (อังกฤษ) และคำแปลไทยด้านล่าง ตรวจสอบคำแปลเทียบกับภาพอย่างละเอียด โดยเน้น:

1. **สลับแถวในตาราง (สำคัญสุด)**: เลขรายการแต่ละข้อ จับคู่กับ "เงื่อนไข/คำอธิบาย" ที่ถูกต้องตรงตามภาพหรือไม่ — ระวังตารางที่เซลล์เงื่อนไขถูกผสาน (merge) ครอบหลายรายการ มักทำให้คำอธิบายเลื่อนผิดแถว
2. **เนื้อหาตกหล่น**: มีข้อความ/ประโยค/รายการใดในภาพที่หายไปจากคำแปลไหม
3. **ความหมายเพี้ยน**: ตัวเลข หน่วย เงื่อนไข (เช่น "เฉพาะสีที่ 1", "25 มม.", "เฉพาะเดินหน้า") ตรงกับต้นฉบับไหม
4. **ห้ามแก้สไตล์/ถ้อยคำที่ถูกอยู่แล้ว** แก้เฉพาะจุดที่ผิดจริงเทียบกับภาพ

กฎสำคัญ:
- "issues" ใส่เฉพาะ "จุดที่ผิดจริงและคุณแก้แล้ว" เท่านั้น — ห้ามใส่รายการที่ถูกต้องอยู่แล้ว ห้ามใส่คำยืนยันว่า "ตรงกับต้นฉบับ"
- ส่วนที่จาง/เลือน/ว่าง อ่านไม่ออก ใส่ `[ส่วนนี้จาง อ่านไม่ออก — ต้องสแกนใหม่]` ครั้งเดียว ห้ามถอดความ ห้ามใส่ตัวอักษร/`&nbsp;` ซ้ำๆ

ตอบเป็น JSON object มี 2 คีย์เท่านั้น:
{
  "issues": ["จุดที่แก้ สั้นๆ เป็นไทย เช่น 'ข้อ 10 เงื่อนไขสลับกับข้อ 11 — แก้แล้ว'", ...],   // ถ้าไม่พบจุดผิดให้เป็น []
  "corrected_markdown": "คำแปลฉบับแก้ไขครบทั้งหน้า (Markdown เดิม แก้เฉพาะจุดผิด คงรูปแบบ/ตาราง/ศัพท์อังกฤษไว้)"
}\
"""

# ---------------------------------------------------------------- vision

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def get_page_text(doc: fitz.Document, idx: int) -> str:
    """ข้อความจาก text layer (ว่าง = หน้าสแกน). คืน '' ถ้าน้อยกว่าเกณฑ์"""
    t = (doc[idx].get_text() or "").strip()
    return t if len(t) >= TEXT_LAYER_MIN else ""


def _with_source(prompt: str, source_text: str | None) -> str:
    """แนบข้อความ text layer เป็น ground-truth ต่อท้าย prompt (โหมดไฮบริด)"""
    if not source_text:
        return prompt
    return (
        prompt
        + "\n\n=== ข้อความจริงจาก text layer ของหน้านี้ (ถูกต้อง 100%) ==="
        + "\nใช้ข้อความนี้เป็นความจริงของ 'ตัวอักษร ตัวเลข ชื่อชิ้นส่วน รหัส หน่วย' — "
        + "ห้ามอ่านเลขจากภาพผิดไปจากนี้ ใช้ภาพเพียงเพื่อดูเลย์เอาต์ ตาราง และรูปประกอบ "
        + "(ข้อความอาจเรียงสลับ ให้จัดลำดับตามภาพ):\n\n"
        + source_text
    )


def render_page_b64(doc: fitz.Document, idx: int, dpi: int, max_side: int) -> str:
    """เรนเดอร์หน้าเป็น JPEG (base64) ลดขนาดถ้าใหญ่เกิน เพื่อคุมต้นทุน token"""
    page = doc[idx]
    zoom = dpi / 72.0
    rect = page.rect
    longest = max(rect.width, rect.height) * zoom
    if longest > max_side:
        zoom *= max_side / longest
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    jpg = pix.tobytes("jpg", jpg_quality=88)
    return base64.b64encode(jpg).decode()


def _sanitize(md: str) -> str:
    """กันอาการ vision ติดลูปพ่นอักขระ/ช่องว่างซ้ำบนหน้าจาง-ว่าง"""
    md = md.replace("&nbsp;", " ")
    md = re.sub(r"[ \t]{3,}", " ", md)               # ช่องว่างยาว -> เดียว
    md = re.sub(r"([.\-_·•*=])\1{4,}", r"\1\1\1", md)  # จุด/ขีดซ้ำยาว -> สั้น
    # บรรทัดเดียวกันซ้ำติดกันเกิน 2 ครั้ง -> เหลือครั้งเดียว
    out, prev, rep = [], None, 0
    for ln in md.split("\n"):
        s = ln.strip()
        if s and s == prev:
            rep += 1
            if rep >= 2:
                continue
        else:
            rep = 0
            prev = s
        out.append(ln)
    return "\n".join(out).strip()


def translate_image(b64: str, source_text: str | None = None) -> tuple[str, int, int]:
    """แปลภาพหน้า -> markdown. ถ้ามี source_text จะใช้โหมดไฮบริด (ตัวอักษรแม่นยำ)"""
    resp = client.chat.completions.create(
        model=MODEL,
        temperature=0.2,
        max_tokens=12000,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": _with_source(PAGE_PROMPT, source_text)},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                ],
            },
        ],
    )
    md = _sanitize(_strip_code_fence((resp.choices[0].message.content or "").strip()))
    u = resp.usage
    return md, u.prompt_tokens, u.completion_tokens


# บรรทัดที่เป็นแค่ "ยืนยันว่าถูก" ไม่ใช่จุดผิดจริง — กรองออกจากรายงาน
_NOISE_RE = re.compile(
    r"ไม่มีจุดผิด|ไม่มีเนื้อหาตกหล่น|ไม่มี(การ)?สลับ|ไม่มีความหมายเพี้ยน|"
    r"ไม่พบจุดผิด|แปลถูกต้อง|จับคู่กับเงื่อนไขถูกต้อง|"
    r"ตรงกับต้นฉบับ(แล้ว)?$|ถูกต้องแล้ว\)?$|ถูกต้องครบถ้วน"
)


def _real_issues(issues: list[str]) -> list[str]:
    seen = [s.strip() for s in issues if s.strip() and not _NOISE_RE.search(s.strip())]
    return list(dict.fromkeys(seen))   # กันบรรทัดซ้ำ คงลำดับเดิม


def verify_image(b64: str, md: str, source_text: str | None = None) -> tuple[str, list[str], int, int]:
    """
    รอบตรวจทาน: ให้ gpt อ่านภาพต้นฉบับซ้ำ เทียบกับคำแปล แล้ว
      - หาจุดที่ "เลขรายการ <-> เงื่อนไข" สลับแถว (ปัญหา merged cell)
      - หาเนื้อหาที่ตกหล่น / แปลเพี้ยน
      - คืน markdown ที่แก้แล้ว + รายการปัญหาที่พบ
    คืน (corrected_md, issues, prompt_tokens, completion_tokens)
    """
    user_text = _with_source(VERIFY_PROMPT, source_text) + "\n\n=== คำแปลที่ต้องตรวจ ===\n" + md
    resp = client.chat.completions.create(
        model=MODEL,
        temperature=0,
        max_tokens=14000,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": VERIFY_SYSTEM},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": user_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                ],
            },
        ],
    )
    u = resp.usage
    raw = resp.choices[0].message.content or "{}"
    try:
        import json
        raw = _strip_code_fence(raw.strip())
        data = json.loads(raw)
        issues = _real_issues([str(x) for x in data.get("issues", [])])
        corrected = _sanitize(_strip_code_fence((data.get("corrected_markdown") or "").strip()))
        if not corrected:          # โมเดลไม่ส่งฉบับแก้ -> ใช้ของเดิม
            corrected = md
    except Exception:              # noqa: BLE001 — JSON พัง ใช้ของเดิม ไม่ให้ทั้ง pipeline ล้ม
        issues = ["[verify] อ่านผล JSON ไม่สำเร็จ — ใช้คำแปลเดิม ควรตรวจหน้านี้ด้วยมือ"]
        corrected = md
    return corrected, issues, u.prompt_tokens, u.completion_tokens


def _strip_code_fence(md: str) -> str:
    """gpt บางทีครอบทั้งหน้าด้วย ```markdown ... ``` — ตัดออก"""
    m = re.match(r"^```[a-zA-Z]*\s*\n(.*)\n```$", md, flags=re.DOTALL)
    return m.group(1).strip() if m else md

# ---------------------------------------------------------------- markdown -> docx

def _set_run_font(run, name: str, size: int | None = None, bold: bool | None = None):
    run.font.name = name
    if size is not None:
        from docx.shared import Pt
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    # ผูกฟอนต์ให้ครอบคลุม ascii + complex-script (ไทย)
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


_BR_RE = re.compile(r"<br\s*/?>", flags=re.IGNORECASE)


def _emit_run(paragraph, text: str, size, bold):
    """เพิ่ม run โดยแปลง <br> เป็นการขึ้นบรรทัดจริงในเซลล์/ย่อหน้า"""
    parts = _BR_RE.split(text)
    for k, part in enumerate(parts):
        r = paragraph.add_run(part)
        _set_run_font(r, THAI_FONT, size, bold)
        if k < len(parts) - 1:
            r.add_break()


def _add_inline(paragraph, text: str, size: int | None = None, base_bold: bool = False):
    """ใส่ข้อความลง paragraph รองรับ **ตัวหนา** และ <br>"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _emit_run(paragraph, text[pos:m.start()], size, base_bold)
        _emit_run(paragraph, m.group(1), size, True)
        pos = m.end()
    if pos < len(text):
        _emit_run(paragraph, text[pos:], size, base_bold)
    if not text:
        _emit_run(paragraph, "", size, base_bold)


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{2,}.*$", line)) and set(line.strip()) <= set("|:- ")


def _split_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_uniform(pix) -> bool:
    """รูปสีเดียวเกือบทั้งหมด (กล่องดำ/แถบตกแต่ง/extraction พลาด) -> ข้าม"""
    s = pix.samples
    if not s:
        return True
    step = max(1, len(s) // 400)
    samp = s[::step]
    return bool(samp) and (max(samp) - min(samp) < 12)


def extract_page_images(doc: fitz.Document, idx: int, out_dir: Path,
                        min_area: int = 20000) -> list[tuple[Path, int, int]]:
    """ดึงรูป raster (รูปถ่าย/สกรีนช็อต) จากหน้า idx เรียงบน->ล่าง
    ข้าม: โลโก้/ไอคอนเล็ก, image mask (bpc<=1), รูปสีเดียว (กล่องดำ)
    คืน [(path, w, h), ...]"""
    out: list[tuple[Path, int, int]] = []
    try:
        page = doc[idx]
        infos = page.get_image_info(xrefs=True)
        meta = {im[0]: im for im in page.get_images(full=True)}   # xref -> tuple
    except Exception:
        return out
    infos = sorted(infos, key=lambda im: (round(im.get("bbox", [0, 0])[1]),
                                          round(im.get("bbox", [0, 0])[0])))
    seen: set[int] = set()
    seq = 0
    for im in infos:
        xref = im.get("xref", 0)
        if not xref or xref in seen:
            continue
        seen.add(xref)
        w, h = int(im.get("width", 0)), int(im.get("height", 0))
        if w * h < min_area:          # ข้ามโลโก้/ไอคอน
            continue
        m = meta.get(xref)
        if m:
            smask, bpc, cs = m[1], m[4], m[5]
            # ข้าม: รูปมี soft-mask (มักเป็นตัวหนังสือหัวเรื่องฝังเป็นภาพโปร่งใส
            # -> ดึงออกมาเป็นกล่องดำ), stencil mask, หรือไม่มี colorspace
            if smask or (bpc and bpc <= 1) or not cs:
                continue
        try:
            pix = fitz.Pixmap(doc, xref)
            if pix.n - pix.alpha >= 4:  # CMYK -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            if _is_uniform(pix):        # กล่องดำ/แถบตกแต่ง -> ข้าม
                continue
            fn = out_dir / f"page_{idx + 1:04d}_img_{seq}.png"
            pix.save(str(fn))
            out.append((fn, pix.width, pix.height))
            seq += 1
        except Exception:
            continue
    return out


def strip_boilerplate(pages_md: list[tuple[int, str]]) -> list[tuple[int, str]]:
    """ตัดหัว/ท้ายกระดาษซ้ำทุกหน้า (ชื่อบริษัท/ชื่อคู่มือที่โผล่ทุกหน้า, เลขหน้า,
    มาร์กเกอร์ footer แบบ *...* ) ออกจากเนื้อหา — ใช้กับงานเก่าที่ cache ไว้ได้เลย"""
    from collections import Counter

    cnt: Counter = Counter()
    for _, md in pages_md:
        for s in {ln.strip() for ln in md.split("\n")}:
            if 15 <= len(s) <= 70 and s[:1] not in "#|->`[" and "ภาพ" not in s \
                    and not s.startswith("**"):
                cnt[s] += 1
    npages = len(pages_md)
    thresh = max(4, int(npages * 0.4))
    boiler = {s for s, c in cnt.items() if c >= thresh}   # โผล่ >=40% ของหน้า = boilerplate

    num = re.compile(r"^\*?\s*\d{1,4}\s*\*?$")            # เลขหน้าโดดๆ (อาจมี *...*)
    star = re.compile(r"^\*[^*].*\*$")                    # บรรทัดห่อด้วย *...* เดี่ยว (footer marker)

    out: list[tuple[int, str]] = []
    for p, md in pages_md:
        kept = []
        for ln in md.split("\n"):
            s = ln.strip()
            if s in boiler:
                continue
            if num.match(s):
                continue
            if star.match(s) and not s.startswith("**"):   # ห่อ * เดี่ยว ไม่ใช่ **bold**
                continue
            kept.append(ln)
        md2 = re.sub(r"\n{3,}", "\n\n", "\n".join(kept)).strip()
        out.append((p, md2))
    return out


def _add_page_numbers(doc):
    """ใส่เลขหน้าแบบ 'หน้า N' เรียงต่อเนื่อง กึ่งกลาง footer (เทาอ่อน)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    GRAY = RGBColor(0xAA, 0xAA, 0xAA)

    def styled(text=None):
        run = p.add_run(text or "")
        _set_run_font(run, THAI_FONT, 10)
        run.font.color.rgb = GRAY
        return run

    try:
        p = doc.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled("หน้า ")
        # โครงสร้าง field {PAGE} ครบ 5 ส่วน — ทุก run จัดสีเทาอ่อนเหมือนกัน
        # รวมถึง "ผลลัพธ์ที่ cache ไว้" (เลขชั่วคราว) เพื่อให้เลขหน้ารับสีเทา ไม่กลับเป็นดำ
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        styled()._r.append(b)
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
        styled()._r.append(it)
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        styled()._r.append(sep)
        styled("1")                       # cached result (จัดสีเทาแล้ว)
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        styled()._r.append(e)
    except Exception:
        pass


def _add_divider(doc):
    """เส้นคั่นจางๆ ระหว่างหน้าต้นฉบับ (ไม่มีตัวเลข)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DDDDDD")
    pbdr.append(bottom)
    pPr.append(pbdr)


def render_markdown_to_docx(pages_md: list[tuple[int, str]], out_docx: Path, title: str,
                            page_images: dict[int, list] | None = None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    page_images = page_images or {}
    doc = Document()
    # ตั้งฟอนต์ + ระยะบรรทัด/ย่อหน้า ให้อ่านสบาย (ภาษาไทยต้องการบรรทัดโปร่ง)
    normal = doc.styles["Normal"]
    normal.font.name = THAI_FONT
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)
    _add_page_numbers(doc)   # footer "หน้า N" เรียงต่อเนื่อง

    # หน้าปก
    doc.add_paragraph()
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(title)
    _set_run_font(r, THAI_FONT, 24, True)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    r = sub.add_run("แปลโดย gpt-4.1 — เอกสารแปลเพื่อใช้งานภายใน Sirivatana Interprint")
    _set_run_font(r, THAI_FONT, 11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()   # ขึ้นหน้าใหม่หลังหน้าปก ครั้งเดียว
    first = True
    for src_page, md in pages_md:
        imgs = page_images.get(src_page, [])
        if not md.strip() and not imgs:
            continue       # หน้าว่างจริง (เช่นมีแต่รูปเวกเตอร์ที่ดึงไม่ได้) -> ข้าม ไม่ทำหน้าว่าง
        # เนื้อหาไหลต่อเนื่อง คั่นหน้าต้นฉบับด้วยเส้นจางๆ (เลขหน้าอยู่ที่ footer แล้ว)
        if not first:
            _add_divider(doc)
        first = False
        _render_page_body(doc, md, imgs)

    doc.save(str(out_docx))


def _render_page_body(doc, md: str, images: list | None = None):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    images = list(images or [])
    img_i = [0]   # ตัวนับรูปที่ใช้ไปแล้วในหน้านี้

    def _next_image():
        """คืนรูปถัดไปของหน้านี้ (path, w, h) แล้วเลื่อนตัวนับ — None ถ้าหมด"""
        if img_i[0] >= len(images):
            return None
        item = images[img_i[0]]
        img_i[0] += 1
        return item

    def _insert_image():
        """แปะรูปถัดไปของหน้านี้ลงเอกสาร (กึ่งกลาง) คืน True ถ้าแปะ"""
        item = _next_image()
        if not item:
            return False
        path, w, h = item
        try:
            # คุมทั้งกว้างและสูง ไม่ให้รูปล้นหน้า + แพ็คเนื้อหาแน่นขึ้น (ลดช่องว่าง)
            max_w, max_h = 5.4, 4.0      # นิ้ว
            width_in = min(max_w, (w / 96.0) if w else max_w)
            if w and h and width_in * (h / w) > max_h:
                width_in = max_h * (w / h)
            doc.add_picture(str(path), width=Inches(width_in))
            p = doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)   # เว้นจากตาราง/เนื้อหาด้านบน ไม่ทะลุกรอบ
            return True
        except Exception:
            return False

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ตาราง: บรรทัด | ... | ตามด้วยบรรทัดคั่น |---|
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            header = _split_row(stripped)
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(_split_row(lines[j].strip()))
                j += 1
            _add_table(doc, header, rows, _next_image)
            i = j
            continue

        if not stripped:
            i += 1
            continue

        # หัวข้อ
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            p = doc.add_paragraph()
            size = {1: 20, 2: 17, 3: 15, 4: 14}.get(level, 13)
            _add_inline(p, m.group(2), size=size, base_bold=True)
            pf = p.paragraph_format
            pf.space_before = Pt(12 if level <= 2 else 8)
            pf.space_after = Pt(4)
            pf.keep_with_next = True   # ไม่ให้หัวข้อค้างท้ายหน้าโดดเดี่ยว
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)   # น้ำเงินเข้ม
            i += 1
            continue

        # เส้นคั่น
        if re.match(r"^([-*_])\1{2,}$", stripped):
            i += 1
            continue

        # ภาพ / placeholder — แปะรูปจริงจากต้นฉบับ (ถ้ามี) + คำแปล caption ใต้รูป
        if (stripped.startswith("[") and "ภาพ" in stripped) or (stripped.startswith("`[") and "ภาพ" in stripped):
            _insert_image()
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("`"))
            _set_run_font(r, THAI_FONT, 11)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x66, 0x99)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # bullet
        bm = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, bm.group(1))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            stripped = stripped.lstrip(">").strip()

        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1

    # รูปที่เหลือ (โมเดลใส่ placeholder ไม่ครบจำนวนรูปจริง) แปะต่อท้ายหน้า
    while img_i[0] < len(images):
        if not _insert_image():
            break


# รูปในช่อง "รูปภาพ/Picture" ของตารางอะไหล่ — โมเดลออกได้หลายแบบ:
#   ![alt](url) , ![alt] (ไม่มี url) , [[alt][image]]
_MD_IMG_RE = re.compile(r"^!\[([^\]]*)\](?:\([^)]*\))?\s*$")
_BRACKET_IMG_RE = re.compile(r"\[\[?([^\]\[]*)\]\s*\[image\]\]?", re.IGNORECASE)


_PIC_CELL_RE = re.compile(r"^\[\s*ภาพ\s*[:：]?\s*(.*?)\s*\]$")   # [ภาพ: alt] ในช่องตาราง


def _cell_image_alt(txt: str) -> str | None:
    """ถ้าเซลล์เป็น 'รูป' คืน alt text (ใช้แสดงถ้าไม่มีรูปจริง); ไม่ใช่รูปคืน None"""
    t = (txt or "").strip()
    m = _MD_IMG_RE.match(t)
    if m:
        return m.group(1).strip() or "(รูป)"
    m3 = _PIC_CELL_RE.match(t)        # [ภาพ: ประแจ 14-17]
    if m3:
        return m3.group(1).strip() or "(รูป)"
    if "[image]" in t.lower():
        m2 = _BRACKET_IMG_RE.search(t)
        return (m2.group(1).strip() if m2 else "") or "(รูป)"
    return None


def _fill_cell(cell, txt, size, bold, next_image):
    """ใส่เนื้อหาลงเซลล์ — ถ้าเป็นรูป ให้ฝังรูปจริงในช่อง (ถ้ามี) ไม่ใช่ก็ใส่ข้อความ"""
    from docx.shared import Inches
    cell.paragraphs[0].clear()
    alt = _cell_image_alt(txt)
    if alt is None:
        _add_inline(cell.paragraphs[0], txt, size=size, base_bold=bold)
        return
    item = next_image() if next_image else None    # บริโภครูป -> เรียงรูปอื่นไม่เลื่อน
    placed = False
    if item:
        path, w, h = item
        try:
            wi = min(1.3, (w / 96.0) if w else 1.3)   # รูปเล็กพอดีช่อง
            if w and h and wi * (h / w) > 1.0:        # คุมความสูง ไม่ให้รูปล้นช่องลงไปทับด้านล่าง
                wi = 1.0 * (w / h)
            cell.paragraphs[0].add_run().add_picture(str(path), width=Inches(wi))
            placed = True
        except Exception:
            placed = False
    if not placed:                                  # ไม่มีรูป -> โชว์ alt text แทน (ไม่ใช่ markdown ดิบ)
        _add_inline(cell.paragraphs[0], alt, size=size, base_bold=bold)


def _add_table(doc, header: list[str], rows: list[list[str]], next_image=None):
    # กันคอลัมน์หาย: บางหน้า header กับ body มีจำนวนคอลัมน์ไม่เท่ากัน
    ncol = max([len(header)] + [len(r) for r in rows]) if rows else len(header)
    header = header + [""] * (ncol - len(header))
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    for c, txt in enumerate(header):
        _fill_cell(table.rows[0].cells[c], txt, 12, True, next_image)
    for row in rows:
        cells = table.add_row().cells
        for c in range(ncol):
            txt = row[c] if c < len(row) else ""
            _fill_cell(cells[c], txt, 12, False, next_image)

# ---------------------------------------------------------------- embed fonts

def _obfuscate_font(guid: str, data: bytes) -> bytes:
    """ทำให้ TTF เป็น .odttf ตามสเปก OOXML — XOR 32 ไบต์แรกด้วย mask จาก GUID (กลับลำดับ)"""
    h = guid.replace("-", "").replace("{", "").replace("}", "")
    mask = bytearray.fromhex(h)
    mask.reverse()
    out = bytearray(data)
    for i in range(min(32, len(out))):
        out[i] ^= mask[i % 16]
    return bytes(out)


def embed_fonts_in_docx(docx_path) -> bool:
    """ฝังฟอนต์ Sarabun ลงไฟล์ .docx เพื่อให้แสดงผลถูกต้องทุกเครื่อง (ไม่ต้องลงฟอนต์
    ไม่มีตัวอักษรกล่อง □ แม้ใน Protected View). คืน True ถ้าฝังสำเร็จ."""
    import uuid
    import zipfile

    variants = [
        ("embedRegular", "Sarabun-Regular.ttf"),
        ("embedBold", "Sarabun-Bold.ttf"),
        ("embedItalic", "Sarabun-Italic.ttf"),
        ("embedBoldItalic", "Sarabun-BoldItalic.ttf"),
    ]
    avail = [(tag, FONT_DIR / fn) for tag, fn in variants if (FONT_DIR / fn).exists()]
    if not avail:
        return False

    docx_path = Path(docx_path)
    R_NS = 'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
    REL_T = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            content = {n: z.read(n) for n in z.namelist()}

        embed_children, rels = [], []
        for i, (tag, path) in enumerate(avail, 1):
            guid = str(uuid.uuid4()).upper()
            content[f"word/fonts/font{i}.odttf"] = _obfuscate_font(guid, path.read_bytes())
            rid = f"rIdSarabun{i}"
            rels.append((rid, f"fonts/font{i}.odttf"))
            embed_children.append(
                f'<w:{tag} r:id="{rid}" w:fontKey="{{{guid}}}" w:subsetted="false"/>'
            )
        font_el = '<w:font w:name="Sarabun">' + "".join(embed_children) + "</w:font>"

        # fontTable.xml — เพิ่มรายการฟอนต์ Sarabun ที่ฝัง
        ft = "word/fontTable.xml"
        if ft in content:
            xml = content[ft].decode("utf-8")
            if "xmlns:r=" not in xml:
                xml = xml.replace("<w:fonts ", "<w:fonts " + R_NS + " ", 1)
            xml = xml.replace("</w:fonts>", font_el + "</w:fonts>", 1)
        else:
            xml = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                   '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                   + R_NS + ">" + font_el + "</w:fonts>")
        content[ft] = xml.encode("utf-8")

        # word/_rels/fontTable.xml.rels
        rels_name = "word/_rels/fontTable.xml.rels"
        rel_lines = "".join(
            f'<Relationship Id="{rid}" Type="{REL_T}" Target="{tgt}"/>' for rid, tgt in rels
        )
        if rels_name in content:
            rl = content[rels_name].decode("utf-8").replace(
                "</Relationships>", rel_lines + "</Relationships>")
        else:
            rl = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                  '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                  + rel_lines + "</Relationships>")
        content[rels_name] = rl.encode("utf-8")

        # [Content_Types].xml — Default สำหรับ .odttf
        ct = content["[Content_Types].xml"].decode("utf-8")
        if 'Extension="odttf"' not in ct:
            ct = ct.replace(
                "</Types>",
                '<Default Extension="odttf" '
                'ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>',
            )
        content["[Content_Types].xml"] = ct.encode("utf-8")

        # settings.xml — เปิดการฝังฟอนต์ (ตำแหน่งต้องถูกตามสคีมา ไม่งั้น Word เมิน!)
        # ลำดับ CT_Settings: ... zoom ... embedTrueTypeFonts ... proofState ... compat
        st = "word/settings.xml"
        if st in content:
            s = content[st].decode("utf-8")
            if "<w:embedTrueTypeFonts" not in s:
                # แทรกก่อน element ตัวแรกที่อยู่ "หลัง" embedTrueTypeFonts ตามสคีมา
                for anchor in ("<w:proofState", "<w:defaultTabStop",
                               "<w:characterSpacingControl", "<w:compat", "</w:settings>"):
                    if anchor in s:
                        s = s.replace(anchor, "<w:embedTrueTypeFonts/>" + anchor, 1)
                        break
            # เลี่ยง Compatibility Mode: 14 (Word 2010) -> 15 (Word 2013+ ทันสมัย)
            s = s.replace('compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="14"',
                          'compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"')
            content[st] = s.encode("utf-8")

        tmp = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
            for n, data in content.items():
                z.writestr(n, data)
        tmp.replace(docx_path)
        return True
    except Exception:
        return False


# ---------------------------------------------------------------- docx -> pdf

def docx_to_pdf(docx_path: Path) -> Path | None:
    soffice = _find_soffice()
    if not soffice:
        return None
    docx_path = Path(docx_path).resolve()   # absolute — as_uri() ต้องการ + soffice ชอบ abs path
    # profile แยกต่อการแปลง (กัน lock เมื่อ LibreOffice ถูกเรียกพร้อมกัน) — as_uri()
    # ให้ file URL ถูกต้องทั้ง Windows (file:///C:/...) และ Linux (file:///app/...)
    profile = docx_path.parent / "_lo_profile"
    try:
        subprocess.run(
            [soffice, "-env:UserInstallation=" + profile.as_uri(),
             "--headless", "--convert-to", "pdf", "--outdir",
             str(docx_path.parent), str(docx_path)],
            check=True, capture_output=True, timeout=600,
        )
    except Exception:  # noqa: BLE001 — ถ้าแปลง PDF ไม่ได้ ก็ยังได้ .docx
        return None
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.exists() else None

# ---------------------------------------------------------------- main

def parse_pages(spec: str | None, total: int) -> list[int]:
    """'1-50' / '3' / '1-10,20-25' -> รายการ index 0-based"""
    if not spec:
        return list(range(total))
    out: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-")
            out.update(range(int(a) - 1, int(b)))
        elif part:
            out.add(int(part) - 1)
    return sorted(p for p in out if 0 <= p < total)


def translate_pdf(
    pdf_path,
    out_dir,
    base: str,
    *,
    pages: str | None = None,
    dpi: int = 230,
    max_side: int = 2600,
    workers: int = 4,
    verify: bool = True,
    use_text: bool = True,
    embed_images: bool = True,
    make_pdf: bool = True,
    progress=None,        # callable(done:int, total:int) — รายงานความคืบหน้า
    log=print,            # callable(msg:str) — ข้อความสถานะ
) -> dict:
    """แปล PDF -> ไฟล์ไทย (.docx [+ .pdf]) คืน dict ผลลัพธ์.
    ใช้ได้ทั้ง CLI และ backend (ผ่าน progress/log callback). idempotent บน cache."""
    import json

    pdf_path = Path(pdf_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    cache_dir = out_dir / "_pages"
    cache_dir.mkdir(exist_ok=True)

    doc = fitz.open(str(pdf_path))
    total_pages = len(doc)
    targets = parse_pages(pages, total_pages)
    n_digital = sum(1 for i in targets if get_page_text(doc, i)) if use_text else 0
    ttag = f"ไฮบริด ({n_digital}/{len(targets)} หน้ามี text)" if use_text else "vision ล้วน"
    log(f"เอกสาร: {pdf_path.name} ({total_pages} หน้า) จะแปล {len(targets)} หน้า "
        f"@ {dpi} DPI x{workers} workers ตรวจทาน:{'เปิด' if verify else 'ปิด'} โหมด:{ttag}")

    tok_in = tok_out = 0

    def work(idx: int):
        n = idx + 1
        raw_c = cache_dir / f"page_{n:04d}.md"
        ver_c = cache_dir / f"page_{n:04d}.verified.md"
        iss_c = cache_dir / f"page_{n:04d}.issues.json"
        pin = pout = 0
        if verify and ver_c.exists() and ver_c.stat().st_size > 0:
            cached_iss = json.loads(iss_c.read_text(encoding="utf-8")) if iss_c.exists() else []
            return idx, ver_c.read_text(encoding="utf-8"), _real_issues(cached_iss), 0, 0, True, ""
        if not verify and raw_c.exists() and raw_c.stat().st_size > 0:
            return idx, raw_c.read_text(encoding="utf-8"), [], 0, 0, True, ""

        src = get_page_text(doc, idx) if use_text else ""
        mode = "ไฮบริด" if src else "vision"
        b64 = render_page_b64(doc, idx, dpi, max_side)
        if raw_c.exists() and raw_c.stat().st_size > 0:
            md = raw_c.read_text(encoding="utf-8")
        else:
            md, p1, p2 = translate_image(b64, src or None)
            pin += p1
            pout += p2
            raw_c.write_text(md, encoding="utf-8")
        issues: list[str] = []
        if verify:
            md, issues, p1, p2 = verify_image(b64, md, src or None)
            pin += p1
            pout += p2
            ver_c.write_text(md, encoding="utf-8")
            iss_c.write_text(json.dumps(issues, ensure_ascii=False, indent=2), encoding="utf-8")
        return idx, md, issues, pin, pout, False, mode

    results: dict[int, str] = {}
    all_issues: dict[int, list[str]] = {}
    done = 0
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futs = {ex.submit(work, idx): idx for idx in targets}
        for fut in as_completed(futs):
            idx, md, issues, pin, pout, cached, mode = fut.result()
            results[idx] = md
            all_issues[idx] = issues
            tok_in += pin
            tok_out += pout
            done += 1
            tag = "cache" if cached else f"{mode} in={pin} out={pout}"
            log(f"  [{done}/{len(targets)}] หน้า {idx + 1} {tag}"
                + (f"  ⚠️ {len(issues)} จุด" if issues else ""))
            if progress:
                try:
                    progress(done, len(targets))
                except Exception:
                    pass

    pages_md = [(idx + 1, results[idx]) for idx in targets]
    pages_md = strip_boilerplate(pages_md)

    page_images: dict[int, list] = {}
    if embed_images:
        img_dir = out_dir / "_images"
        img_dir.mkdir(exist_ok=True)
        n_img = 0
        for idx in targets:
            imgs = extract_page_images(doc, idx, img_dir)
            if imgs:
                page_images[idx + 1] = imgs
                n_img += len(imgs)
        log(f"ดึงรูปจากต้นฉบับ: {n_img} รูป ใน {len(page_images)} หน้า")

    title = f"คู่มือ (แปลไทย) — {base}"
    out_docx = out_dir / f"{base}_แปลไทย.docx"
    combined_md = out_dir / f"{base}_แปลไทย.md"
    combined_md.write_text(
        "\n\n".join(f"<!-- หน้า {p} -->\n\n{m}" for p, m in pages_md), encoding="utf-8"
    )
    log("กำลังสร้าง Word ...")
    render_markdown_to_docx(pages_md, out_docx, title, page_images)
    if embed_fonts_in_docx(out_docx):       # ฝัง Sarabun -> แสดงถูกทุกเครื่อง ไม่มี □
        log("ฝังฟอนต์ Sarabun ลงไฟล์แล้ว")

    pdf_out = None
    if make_pdf:
        log("กำลังสร้าง PDF (LibreOffice) ...")
        pdf_out = docx_to_pdf(out_docx)

    flagged = [(idx + 1, all_issues[idx]) for idx in targets if all_issues.get(idx)]
    review = None
    if verify:
        review = out_dir / f"{base}_ตรวจทาน.md"
        lines = [f"# รายงานตรวจทานการแปล — {base}", ""]
        if flagged:
            lines.append(f"พบจุดที่แก้/ควรตรวจซ้ำ ใน {len(flagged)} หน้า (จาก {len(targets)} หน้า):\n")
            for p, issues in flagged:
                lines.append(f"## หน้า {p}")
                lines.extend(f"- {it}" for it in issues)
                lines.append("")
        else:
            lines.append("ไม่พบจุดผิดจากรอบตรวจทาน ✅ (ยังแนะนำให้คนสุ่มตรวจหน้าตารางสำคัญ)")
        review.write_text("\n".join(lines), encoding="utf-8")

    cost = tok_in / 1e6 * PRICE_IN + tok_out / 1e6 * PRICE_OUT
    doc.close()
    return {
        "out_dir": str(out_dir),
        "docx": str(out_docx),
        "pdf": str(pdf_out) if pdf_out else None,
        "md": str(combined_md),
        "review": str(review) if review else None,
        "n_pages": len(targets),
        "n_flagged": len(flagged),
        "tokens_in": tok_in,
        "tokens_out": tok_out,
        "cost_usd": round(cost, 4),
    }


def main():
    ap = argparse.ArgumentParser(description="แปลคู่มือเครื่องจักร EN->ไทย (gpt-4.1 vision)")
    ap.add_argument("pdf", help="path ไฟล์ PDF ต้นฉบับ (สแกนได้)")
    ap.add_argument("--out", help="ชื่อไฟล์ผลลัพธ์ (ไม่ต้องใส่นามสกุล)", default=None)
    ap.add_argument("--pages", help="ช่วงหน้า เช่น 1-50 หรือ 1-10,20-25 (ดีฟอลต์ทั้งเล่ม)")
    ap.add_argument("--dpi", type=int, default=230)
    ap.add_argument("--max-side", type=int, default=2600)
    ap.add_argument("--workers", type=int, default=4)
    ap.add_argument("--no-verify", action="store_true")
    ap.add_argument("--no-text-layer", action="store_true")
    ap.add_argument("--no-images", action="store_true")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        sys.exit(f"ไม่พบไฟล์: {pdf_path}")
    base = args.out or pdf_path.stem
    out_dir = pdf_path.parent / f"{base}_แปลไทย"

    res = translate_pdf(
        pdf_path, out_dir, base,
        pages=args.pages, dpi=args.dpi, max_side=args.max_side, workers=args.workers,
        verify=not args.no_verify, use_text=not args.no_text_layer,
        embed_images=not args.no_images, make_pdf=not args.no_pdf,
    )
    print("\n===== สรุป =====")
    print(f"tokens in={res['tokens_in']} out={res['tokens_out']}")
    print(f"ต้นทุน ~= ${res['cost_usd']:.3f}  (~{res['cost_usd'] * USD_TO_THB:.2f} บาท)")
    print(f"docx: {res['docx']}")
    if res["pdf"]:
        print(f"pdf:  {res['pdf']}")
    if res["review"]:
        print(f"ตรวจทาน: {res['review']}  ({res['n_flagged']} หน้ามีจุด)")
    print(f"ไฟล์อยู่ที่: {res['out_dir']}")


if __name__ == "__main__":
    main()
