"""Convert a chat answer (Markdown) into a Word .docx.

Used by POST /chat/export-docx. Handles the elements that show up in answers:
headings, paragraphs, **bold**, bullet/numbered lists, and GFM tables. Not a
full Markdown engine — just the common cases — but produces a clean, editable
Word document.
"""

import io
import re

from docx import Document
from docx.shared import Pt

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_SEP_CELL = re.compile(r"^:?-{1,}:?$")
_HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = [c for c in _split_row(line) if c != ""]
    return len(cells) > 0 and all(_SEP_CELL.match(c) for c in cells)


def _add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, turning **...** into bold runs."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def export_markdown_to_docx(content_md: str, title: str | None = None) -> bytes:
    doc = Document()

    # Tahoma renders Thai cleanly across Office versions.
    normal = doc.styles["Normal"]
    normal.font.name = "Tahoma"
    normal.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    lines = (content_md or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # GFM table: header row + |---| separator + data rows.
        if (
            "|" in line
            and stripped
            and i + 1 < len(lines)
            and _is_separator(lines[i + 1])
        ):
            header = _split_row(line)
            rows = [header]
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                rows.append(_split_row(lines[j]))
                j += 1
            ncol = len(header)
            table = doc.add_table(rows=0, cols=ncol)
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                cells = table.add_row().cells
                for c in range(ncol):
                    val = row[c] if c < len(row) else ""
                    para = cells[c].paragraphs[0]
                    _add_inline(para, val)
                    if r == 0:
                        for run in para.runs:
                            run.bold = True
            i = j
            continue

        if not stripped:
            i += 1
            continue

        mh = _HEADING_RE.match(stripped)
        if mh:
            doc.add_heading(mh.group(2), level=min(len(mh.group(1)), 4))
            i += 1
            continue

        if _HR_RE.match(stripped):
            i += 1
            continue

        mb = _BULLET_RE.match(stripped)
        if mb:
            _add_inline(doc.add_paragraph(style="List Bullet"), mb.group(1))
            i += 1
            continue

        mn = _NUMBERED_RE.match(stripped)
        if mn:
            _add_inline(doc.add_paragraph(style="List Number"), mn.group(1))
            i += 1
            continue

        _add_inline(doc.add_paragraph(), stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
